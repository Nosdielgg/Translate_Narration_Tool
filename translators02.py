# ===============================================
# === FUNÇÕES DE TRADUÇÃO (continuação) ===
# ===============================================

import os
from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Pt
from .file_handlers import read_file_text, read_docx_text, read_srt, apply_paragraph_formatting
from .utils import safe_filename
from .translators import translate_text_in_segments

def translate_file(input_path, output_folder, target_language_key, status_callback=None):
    """Traduz arquivos (SRT, DOCX, TXT) e salva como DOCX."""
    if status_callback:
        status_callback("Lendo arquivo de entrada...", "blue")
    
    from .constants import LANGUAGES
    
    target_lang_code = LANGUAGES.get(target_language_key, "en")
    base_name, ext = os.path.splitext(os.path.basename(input_path))
    ext = ext.lower()
    
    original_segments_list = []
    
    try:
        if ext == ".srt":
            text_blocks = read_srt(input_path)
            if not text_blocks:
                raise ValueError("Não foi possível ler blocos de texto do arquivo SRT.")
            full_text = '\n\n'.join(text_blocks)
            original_segments_list = text_blocks
            
        elif ext == ".docx":
            full_text = read_docx_text(input_path)
            original_segments_list = [p.strip() for p in full_text.split('\n') if p.strip()]
            
        elif ext == ".txt":
            full_text = read_file_text(input_path)
            original_segments_list = [p.strip() for p in full_text.split('\n') if p.strip()]
        
        else:
            raise ValueError("Formato de arquivo não suportado. Use .srt, .docx ou .txt.")
        
    except Exception as e:
        if status_callback:
            status_callback("Erro de leitura.", "red")
        raise Exception(f"Falha ao ler o arquivo: {e}")

    if status_callback:
        status_callback(f"Iniciando tradução para {target_language_key}...", "blue")
    
    translated_texts = translate_text_in_segments(full_text, target_lang_code, status_callback)
    
    if not translated_texts:
        if status_callback:
            status_callback("Tradução falhou.", "red")
        return False

    # Salvar como DOCX
    try:
        doc = Document()
        
        # Título
        title_para = doc.add_paragraph(f"Translation Report ({base_name})")
        title_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        title_para.runs[0].font.size = Pt(14)
        title_para.runs[0].font.bold = True
        
        doc.add_paragraph(f"Target Language: {target_language_key}")
        doc.add_paragraph(f"Original Format: {ext.upper()}")
        doc.add_paragraph("-" * 80)
        
        # Corpo do documento (modo bilíngue)
        translated_segments_list = [t.strip() for t in '\n'.join(translated_texts).split('\n') if t.strip()]
        
        if len(original_segments_list) == len(translated_segments_list):
            for i, (orig_text, trans_text) in enumerate(zip(original_segments_list, translated_segments_list)):
                # Texto original
                para_orig = doc.add_paragraph()
                para_orig.add_run(f"Original Segment {i+1}: ").bold = True
                para_orig.add_run(orig_text)
                apply_paragraph_formatting(para_orig)
                
                # Tradução
                para_trans = doc.add_paragraph()
                para_trans.add_run(f"Translation: ").bold = True
                para_trans.add_run(trans_text)
                apply_paragraph_formatting(para_trans)
                
                doc.add_paragraph()  # Espaçamento
        else:
            # Fallback para apenas tradução
            doc.add_paragraph("WARNING: Segment mismatch detected. Saving only translation.")
            doc.add_paragraph("-" * 80)
            
            full_translated_text = '\n\n'.join(translated_texts)
            for trans_paragraph in full_translated_text.split('\n\n'):
                if trans_paragraph.strip():
                    para_trans = doc.add_paragraph()
                    para_trans.add_run(trans_paragraph.strip())
                    apply_paragraph_formatting(para_trans)

        # Lógica de salvamento
        os.makedirs(output_folder, exist_ok=True)
        safe_name = safe_filename(base_name)
        
        base_path = os.path.join(output_folder, f"{safe_name}_Translated_{target_lang_code}.docx")
        final_path = base_path
        
        counter = 1
        while os.path.exists(final_path):
            final_path = os.path.join(output_folder, f"{safe_name}_Translated_{target_lang_code}({counter}).docx")
            counter += 1

        doc.save(final_path)
        
        if status_callback:
            status_callback("Tradução e geração de DOCX concluídas!", "green")
        return final_path
    
    except Exception as e:
        if status_callback:
            status_callback("Erro DOCX.", "red")
        raise Exception(f"Falha ao gerar documento de saída: {e}")
